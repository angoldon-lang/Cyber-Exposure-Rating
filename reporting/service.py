"""Generazione dei report: HTML, PDF, Word, JSON e CSV."""
from __future__ import annotations

import csv
import hashlib
import io
import json
from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from jinja2 import Environment, FileSystemLoader, select_autoescape

from app.core.config import settings
from app.core.logging import get_logger
from reporting.context import (
    CONFIDENCE_LABEL_IT,
    EFFORT_LABEL_IT,
    OWNERSHIP_LABEL_IT,
    PRIORITY_LABEL_IT,
    SEVERITY_LABEL_IT,
    ReportContext,
)

logger = get_logger(__name__)

TEMPLATE_DIR = Path(__file__).resolve().parent / "templates"


@dataclass
class GeneratedReport:
    format: str
    content: bytes
    filename: str

    @property
    def sha256(self) -> str:
        return hashlib.sha256(self.content).hexdigest()

    @property
    def size(self) -> int:
        return len(self.content)


def _environment() -> Environment:
    # `autoescape` attivo: nessun contenuto raccolto da Internet puo' iniettare
    # markup nel report.
    env = Environment(
        loader=FileSystemLoader(str(TEMPLATE_DIR)),
        autoescape=select_autoescape(["html", "xml", "j2"]),
        trim_blocks=True, lstrip_blocks=True)
    return env


def _stylesheet() -> str:
    return (TEMPLATE_DIR / "base.css").read_text(encoding="utf-8")


def render_html(context: ReportContext, template_name: str) -> str:
    template = _environment().get_template(template_name)
    return template.render(**context.as_dict(), stylesheet=_stylesheet())


def _slug(value: str) -> str:
    keep = [c.lower() if c.isalnum() else "-" for c in value]
    return "".join(keep).strip("-")[:60] or "report"


# ---------------------------------------------------------------------------
def generate_html(context: ReportContext, *, include_technical: bool = True) -> GeneratedReport:
    parts = [render_html(context, "executive.html.j2")]
    if include_technical:
        parts.append(render_html(context, "technical.html.j2"))
    content = "\n".join(parts).encode("utf-8")
    name = f"defenix-exposure-rating-{_slug(context.company_name)}-{context.generated_at:%Y%m%d}"
    return GeneratedReport("html", content, f"{name}.html")


def generate_pdf(context: ReportContext, *, include_technical: bool = True) -> GeneratedReport:
    """PDF via WeasyPrint. Se la libreria non e' disponibile viene sollevata
    un'eccezione esplicita: il chiamante degrada su HTML."""
    from weasyprint import HTML  # import locale: dipendenze di sistema pesanti

    documents = [HTML(string=render_html(context, "executive.html.j2")).render()]
    if include_technical:
        documents.append(HTML(string=render_html(context, "technical.html.j2")).render())

    pages = [page for document in documents for page in document.pages]
    pdf_bytes = documents[0].copy(pages).write_pdf()
    name = f"defenix-exposure-rating-{_slug(context.company_name)}-{context.generated_at:%Y%m%d}"
    return GeneratedReport("pdf", pdf_bytes, f"{name}.pdf")


def generate_docx(context: ReportContext, *, include_technical: bool = True) -> GeneratedReport:
    """Report Word tramite python-docx."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    data = context.as_dict()
    document = Document()

    # --- copertina ---
    title = document.add_heading(f"{context.brand['name']} Exposure Rating", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph(
        "Valutazione dell'esposizione cyber osservabile dall'esterno")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph()
    heading = document.add_paragraph()
    run = heading.add_run(context.company_name)
    run.bold = True
    run.font.size = Pt(18)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Data: {context.generated_at:%d/%m/%Y}\n"
                 f"Profilo di scansione: {context.profile_label}")

    document.add_paragraph()
    rating = document.add_paragraph()
    rating.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if context.is_provisional:
        run = rating.add_run("Valutazione provvisoria")
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0xC2, 0x41, 0x0C)
        document.add_paragraph(context.provisional_notice or "").alignment = \
            WD_ALIGN_PARAGRAPH.CENTER
    else:
        run = rating.add_run(f"{context.overall_score:.0f}/100 - Classe {context.rating_class}")
        run.bold = True
        run.font.size = Pt(24)
        document.add_paragraph(context.rating_label).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(
        f"Affidabilita' della rilevazione: {context.confidence_value:.0f}% "
        f"({context.confidence_label})").alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph()
    disclaimer = document.add_paragraph()
    disclaimer.add_run(context.disclaimer).italic = True
    document.add_page_break()

    # --- perimetro ---
    document.add_heading("Perimetro della valutazione", level=1)
    scope_table = document.add_table(rows=0, cols=2)
    scope_table.style = "Light Grid Accent 1"
    for label, value in (
        ("Profilo di scansione", context.profile_label),
        ("Domini analizzati", ", ".join(context.scope.get("domains", [])) or "nessuno"),
        ("Domini verificati", ", ".join(context.scope.get("verified_domains", []))
         or "nessun dominio verificato"),
        ("Indirizzi IP autorizzati", ", ".join(context.scope.get("ip_addresses", [])) or "-"),
        ("Reti autorizzate", ", ".join(context.scope.get("network_ranges", [])) or "-"),
        ("Esclusioni", ", ".join(context.scope.get("excluded", [])) or "nessuna"),
    ):
        row = scope_table.add_row().cells
        row[0].text = label
        row[1].text = value

    # --- rating tematici ---
    document.add_heading("Rating per area tematica", level=1)
    table = document.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    header = table.rows[0].cells
    for index, label in enumerate(("Area", "Peso", "Punteggio", "Rilievi")):
        header[index].text = label
    for category in context.categories:
        row = table.add_row().cells
        row[0].text = str(category.get("label_it", category.get("key")))
        row[1].text = f"{float(category.get('weight', 0)) * 100:.0f}%"
        row[2].text = f"{float(category.get('score', 0)):.0f}/100"
        row[3].text = str(category.get("finding_count", 0))

    # --- rischi principali ---
    document.add_heading("Principali rischi rilevati", level=1)
    if context.top_risks:
        risks = document.add_table(rows=1, cols=4)
        risks.style = "Light Grid Accent 1"
        header = risks.rows[0].cells
        for index, label in enumerate(("Rif.", "Severita'", "Rilievo", "Attendibilita'")):
            header[index].text = label
        for risk in context.top_risks:
            row = risks.add_row().cells
            row[0].text = str(risk.get("reference_code", ""))
            row[1].text = SEVERITY_LABEL_IT.get(str(risk.get("severity")), "")
            row[2].text = str(risk.get("title", ""))
            row[3].text = CONFIDENCE_LABEL_IT.get(str(risk.get("confidence_class")), "")
    else:
        document.add_paragraph("Nessun rilievo significativo nel perimetro analizzato. "
                               "L'assenza di rilievi non costituisce prova di sicurezza.")

    # --- piano di remediation ---
    document.add_heading("Priorita' di intervento", level=1)
    for index, item in enumerate(context.remediation_plan[:10], start=1):
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(f"{item['title_it']} ").bold = True
        paragraph.add_run(
            f"(priorita': {PRIORITY_LABEL_IT.get(item['priority'], item['priority'])}, "
            f"impegno: {EFFORT_LABEL_IT.get(item['effort'], item['effort'])})\n")
        paragraph.add_run(f"Azione immediata: {item['immediate_action_it']}").italic = True

    if context.quick_wins:
        document.add_heading("Interventi rapidi ad alto beneficio", level=2)
        for item in context.quick_wins:
            document.add_paragraph(f"{item['title_it']}: {item['immediate_action_it']}",
                                   style="List Bullet")

    # --- confronto ---
    document.add_heading("Confronto con la scansione precedente", level=1)
    if context.comparison and context.comparison.get("previous_score") is not None:
        document.add_paragraph(str(context.comparison.get("summary_it", "")))
    else:
        document.add_paragraph("Prima valutazione: nessun confronto disponibile.")

    # --- limiti ---
    document.add_heading("Limiti della valutazione", level=1)
    for limit in context.limits:
        document.add_paragraph(limit, style="List Bullet")

    # --- allegato tecnico ---
    if include_technical:
        document.add_page_break()
        document.add_heading("Allegato tecnico", level=1)
        document.add_paragraph(
            "Evidenze sanitizzate. Non sono riportati credenziali, token, cookie, contenuti "
            "integrali di leak, istruzioni di sfruttamento ne' payload offensivi.")
        for finding in context.findings:
            document.add_heading(
                f"{finding.get('reference_code')} - {finding.get('title')}", level=2)
            details = document.add_paragraph()
            details.add_run(
                f"Severita': {SEVERITY_LABEL_IT.get(str(finding.get('severity')), '')} | "
                f"Attendibilita': {CONFIDENCE_LABEL_IT.get(str(finding.get('confidence_class')), '')} | "
                f"Asset: {finding.get('asset_display') or finding.get('detail') or 'n/d'} "
                f"({OWNERSHIP_LABEL_IT.get(str(finding.get('ownership_status')), '')})").italic = True
            document.add_paragraph(str(finding.get("description", "")))
            if finding.get("cve_id"):
                document.add_paragraph(
                    f"CVE: {finding['cve_id']} | CVSS: {finding.get('cvss_score', 'n/d')} | "
                    f"EPSS: {finding.get('epss_score', 'n/d')} | "
                    f"CISA KEV: {'si' if finding.get('cisa_kev') else 'no'}")
            remediation = finding.get("remediation")
            if remediation:
                document.add_paragraph(f"Remediation: {remediation['title_it']}")
                document.add_paragraph(f"Azione immediata: {remediation['immediate_action_it']}")
                document.add_paragraph(f"Verifica: {remediation['verification_it']}")

    document.add_paragraph()
    document.add_paragraph(context.disclaimer).italic = True

    buffer = io.BytesIO()
    document.save(buffer)
    name = f"defenix-exposure-rating-{_slug(context.company_name)}-{context.generated_at:%Y%m%d}"
    return GeneratedReport("docx", buffer.getvalue(), f"{name}.docx")


def generate_json(context: ReportContext) -> GeneratedReport:
    payload = context.as_dict()
    payload.pop("severity_label", None)
    payload.pop("confidence_label_map", None)
    payload.pop("ownership_label", None)
    payload.pop("effort_label", None)
    payload.pop("priority_label", None)
    payload["generated_at"] = context.generated_at.isoformat()
    payload["schema_version"] = "1.0.0"
    content = json.dumps(payload, ensure_ascii=False, indent=2, default=str).encode("utf-8")
    name = f"defenix-exposure-rating-{_slug(context.company_name)}-{context.generated_at:%Y%m%d}"
    return GeneratedReport("json", content, f"{name}.json")


CSV_COLUMNS = [
    "reference_code", "severity", "category", "finding_type", "title", "confidence_class",
    "ownership_status", "asset", "workflow_state", "analyst_validation", "excluded_from_rating",
    "cve_id", "cvss_score", "epss_score", "cisa_kev", "internet_facing",
    "first_seen_at", "last_seen_at", "event_date", "applied_deduction",
    "remediation_id", "remediation_title", "remediation_priority", "remediation_effort",
    "sources",
]


def generate_csv(context: ReportContext) -> GeneratedReport:
    buffer = io.StringIO()
    writer = csv.DictWriter(buffer, fieldnames=CSV_COLUMNS, extrasaction="ignore",
                            delimiter=";", quoting=csv.QUOTE_MINIMAL)
    writer.writeheader()
    for finding in context.findings:
        remediation = finding.get("remediation") or {}
        writer.writerow({
            "reference_code": finding.get("reference_code"),
            "severity": finding.get("severity"),
            "category": finding.get("category"),
            "finding_type": finding.get("finding_type"),
            "title": finding.get("title"),
            "confidence_class": finding.get("confidence_class"),
            "ownership_status": finding.get("ownership_status"),
            "asset": finding.get("asset_display") or finding.get("detail") or "",
            "workflow_state": finding.get("workflow_state"),
            "analyst_validation": finding.get("analyst_validation"),
            "excluded_from_rating": finding.get("excluded_from_rating"),
            "cve_id": finding.get("cve_id") or "",
            "cvss_score": finding.get("cvss_score") if finding.get("cvss_score") is not None else "",
            "epss_score": finding.get("epss_score") if finding.get("epss_score") is not None else "",
            "cisa_kev": finding.get("cisa_kev"),
            "internet_facing": finding.get("internet_facing"),
            "first_seen_at": finding.get("first_seen_at"),
            "last_seen_at": finding.get("last_seen_at"),
            "event_date": finding.get("event_date") or "",
            "applied_deduction": finding.get("applied_deduction", 0),
            "remediation_id": remediation.get("catalog_id") or remediation.get("id", ""),
            "remediation_title": remediation.get("title_it", ""),
            "remediation_priority": remediation.get("priority", ""),
            "remediation_effort": remediation.get("effort", ""),
            "sources": ", ".join(finding.get("sources") or []),
        })
    # BOM UTF-8: Excel in ambiente italiano apre correttamente il file.
    content = buffer.getvalue().encode("utf-8-sig")
    name = f"defenix-findings-{_slug(context.company_name)}-{context.generated_at:%Y%m%d}"
    return GeneratedReport("csv", content, f"{name}.csv")


GENERATORS = {
    "html": generate_html,
    "pdf": generate_pdf,
    "docx": generate_docx,
    "json": generate_json,
    "csv": generate_csv,
}


def generate(context: ReportContext, formats: list[str], *,
             include_technical: bool = True) -> list[GeneratedReport]:
    """Genera i formati richiesti. Un formato non disponibile non blocca gli altri."""
    produced: list[GeneratedReport] = []
    for fmt in formats:
        generator = GENERATORS.get(fmt)
        if generator is None:
            logger.warning("report_format_unknown", format=fmt)
            continue
        try:
            if fmt in {"html", "pdf", "docx"}:
                produced.append(generator(context, include_technical=include_technical))
            else:
                produced.append(generator(context))
        except Exception as exc:  # noqa: BLE001
            logger.error("report_generation_failed", format=fmt, error=str(exc))
            if fmt == "pdf":
                # Degrada su HTML: il contenuto resta consultabile.
                logger.warning("pdf_fallback_to_html")
                produced.append(generate_html(context, include_technical=include_technical))
    return produced


def store(report: GeneratedReport, scan_id: str, version: int,
          base_path: Path | None = None) -> str:
    base = (base_path or settings.report_storage_path) / str(scan_id) / f"v{version}"
    base.mkdir(parents=True, exist_ok=True)
    path = base / report.filename
    path.write_bytes(report.content)
    path.chmod(0o600)
    return str(path)
